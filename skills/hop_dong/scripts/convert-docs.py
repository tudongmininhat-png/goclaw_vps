import argparse
import csv
import difflib
import json
import os
import re
import sys
import tempfile
from copy import deepcopy
from datetime import datetime

from docx import Document
from docx.table import _Row
from docxtpl import DocxTemplate

from quotation_content import (
    build_quotation_content,
    build_quotation_placeholder_content,
    CONTRACT_PAYMENT_TERMS,
    CONTRACT_DELIVERY_LOCATION,
    CONTRACT_SELLER,
)


REQUIRED_FIELDS = {
    "contract": ["TEN_DOI_TAC", "dia_chi", "NGUOI_DAI_DIEN"],
    "quotation": ["TEN_DOI_TAC", "items"],
}
TAG_RE = re.compile(r"{%\s*(.*?)\s*%}")
NUMBER_RE = re.compile(r"[^0-9,.\-+]")
LOOP_OPEN_RE = re.compile(r"{%\s*(?:tr\s+)?(for\s+.+?)\s*%}")
LOOP_CLOSE_RE = re.compile(r"{%\s*(?:tr\s+)?endfor\s*%}")


def sanitize_filename(filename):
    if not filename:
        return "Unknown"
    safe = re.sub(r'[\\/*?:"<>|]', "", str(filename))
    return safe.strip() or "Unknown"


def normalize_model(model_code):
    if not model_code:
        return ""
    return re.sub(r"[^a-zA-Z0-9]", "", str(model_code)).lower()


def force_float(value):
    if isinstance(value, (int, float)):
        return float(value)
    if value is None:
        return 0.0

    clean_str = NUMBER_RE.sub("", str(value).strip()).replace(" ", "")
    if not clean_str:
        return 0.0

    if "," in clean_str and "." in clean_str:
        # Dấu phân cách xuất hiện sau cùng được xem là phần thập phân.
        if clean_str.rfind(",") > clean_str.rfind("."):
            clean_str = clean_str.replace(".", "").replace(",", ".")
        else:
            clean_str = clean_str.replace(",", "")
    elif "," in clean_str:
        parts = clean_str.split(",")
        if len(parts) == 2 and len(parts[1]) <= 2:
            clean_str = f"{parts[0]}.{parts[1]}"
        else:
            clean_str = clean_str.replace(",", "")
    elif "." in clean_str:
        parts = clean_str.split(".")
        if len(parts) > 2:
            clean_str = clean_str.replace(".", "")
        elif len(parts) == 2 and len(parts[1]) > 2:
            clean_str = clean_str.replace(".", "")

    try:
        return float(clean_str)
    except ValueError:
        return 0.0


def format_currency(value):
    try:
        return "{:,.0f}".format(float(value))
    except (TypeError, ValueError):
        return "0"


def get_all_models(csv_path):
    if not os.path.exists(csv_path):
        return []
    models = []
    try:
        with open(csv_path, mode="r", encoding="utf-8-sig") as csv_file:
            reader = csv.DictReader(csv_file)
            for row in reader:
                code = str(row.get("Mã Model", "")).strip()
                if code:
                    models.append(code)
    except Exception:
        pass
    return models


def find_closest_models(query, csv_path, n=3):
    all_models = get_all_models(csv_path)
    if not all_models:
        return []

    # Map normalized models to their original codes
    norm_map = {normalize_model(m): m for m in all_models}
    target_norm = normalize_model(query)

    # Check for exact normalized match first
    if target_norm in norm_map:
        return [norm_map[target_norm]]

    # Otherwise do fuzzy matching on normalized keys
    best_norm_matches = difflib.get_close_matches(target_norm, list(norm_map.keys()), n=n, cutoff=0.3)
    return [norm_map[m] for m in best_norm_matches]



def _fmt_dung_tich(raw: str) -> str:
    """Chuyển '1526.0' → '1,526 L', '520' → '520 L', '' → ''"""
    if not raw:
        return ""
    try:
        val = float(raw)
        int_val = int(val) if val == int(val) else val
        formatted = f"{int_val:,}".replace(",", ".")
        return f"{formatted} L"
    except (ValueError, TypeError):
        return str(raw)


def lookup_product(model_code, csv_path):
    if not os.path.exists(csv_path):
        return None

    target = normalize_model(model_code)
    if not target:
        return None

    try:
        with open(csv_path, mode="r", encoding="utf-8-sig") as csv_file:
            reader = csv.DictReader(csv_file)
            for row in reader:
                if normalize_model(row.get("Mã Model", "")) == target:
                    return {
                        "san_pham": row.get("Tên Model", "") or "",
                        "dung_tich": _fmt_dung_tich(row.get("Dung tích (Lít)", "") or ""),
                        "kich_thuoc": row.get("Kích thước (w x D x H) (mm)", "") or "",
                    }
    except Exception:
        pass
    return None


def _normalize_tag(match_obj):
    inner = re.sub(r"\s+", " ", match_obj.group(1).strip())
    if inner.startswith("tr "):
        # DocxTemplate yêu cầu thẻ phải là {%tr ... %} (không có khoảng trắng sau {%).
        return "{%" + inner + " %}"
    return "{% " + inner + " %}"


def _rewrite_paragraph_text(paragraph, text):
    if not paragraph:
        return

    if not paragraph.runs:
        paragraph.add_run(text)
        return

    first_run = paragraph.runs[0]
    style = first_run.style
    bold = first_run.bold
    font_name = first_run.font.name
    font_size = first_run.font.size

    paragraph_element = paragraph._element
    for run in list(paragraph.runs):
        paragraph_element.remove(run._element)

    new_run = paragraph.add_run(text)
    if style:
        new_run.style = style
    if bold is not None:
        new_run.bold = bold
    if font_name:
        new_run.font.name = font_name
    if font_size:
        new_run.font.size = font_size


def _clean_p_final(paragraph):
    if not paragraph or "{%" not in paragraph.text:
        return

    text = re.sub(r"\s+", " ", paragraph.text).strip()
    text = TAG_RE.sub(_normalize_tag, text)
    _rewrite_paragraph_text(paragraph, text)


def _set_cell_text(cell, text):
    if not cell.paragraphs:
        cell.text = text
        return

    _rewrite_paragraph_text(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        _rewrite_paragraph_text(paragraph, "")


def _normalize_inline_loop_row(table, row):
    row_text = "".join(cell.text for cell in row.cells)
    open_match = LOOP_OPEN_RE.search(row_text)
    close_match = LOOP_CLOSE_RE.search(row_text)
    if not open_match or not close_match:
        return False

    open_tag = "{%tr " + re.sub(r"\s+", " ", open_match.group(1).strip()) + " %}"
    close_tag = "{%tr endfor %}"
    open_tr = deepcopy(row._tr)
    end_tr = deepcopy(row._tr)
    row._tr.addprevious(open_tr)
    row._tr.addnext(end_tr)

    open_row = _Row(open_tr, table)
    end_row = _Row(end_tr, table)

    for cell in open_row.cells:
        _set_cell_text(cell, open_tag)

    for cell in end_row.cells:
        _set_cell_text(cell, close_tag)

    for cell in row.cells:
        if not cell.paragraphs:
            continue
        cleaned_text = cell.text
        cleaned_text = LOOP_OPEN_RE.sub("", cleaned_text)
        cleaned_text = LOOP_CLOSE_RE.sub("", cleaned_text)
        cleaned_text = re.sub(r"\s+", " ", cleaned_text).strip()
        _set_cell_text(cell, cleaned_text)

    return True


def _normalize_quotation_template(document):
    placeholder_content = build_quotation_placeholder_content()
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith("Kính gửi:") and "{{TEN_DOI_TAC}}" not in text:
            _rewrite_paragraph_text(paragraph, placeholder_content["greeting"])
        elif "giao tận nơi tại" in text and "{{dia_chi}}" not in text:
            _rewrite_paragraph_text(paragraph, placeholder_content["notes"][2])
        elif text.startswith("Tp. HCM, Ngày") and "{{NGAY_HIEN_TAI}}" not in text:
            _rewrite_paragraph_text(paragraph, placeholder_content["date_line"])


def _normalize_contract_template(document):
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        
        # Đồng bộ thông tin BÊN B (BÊN BÁN) từ config
        if text.startswith("BÊN B (BÊN BÁN):"):
            _rewrite_paragraph_text(paragraph, CONTRACT_SELLER["line1_name"])
        elif text.startswith("Tài khoản số"):
            _rewrite_paragraph_text(paragraph, CONTRACT_SELLER["line2_account"])
        elif text.startswith("Địa chỉ") and ":" in text and "BÊN B" in text:
            # Lưu ý: "Địa chỉ" xuất hiện ở nhiều nơi, ta cần phân biệt địa chỉ Bên B
            # Ở đây đơn giản hóa bằng cách check keyword hoặc thứ tự, 
            # nhưng tốt nhất là check prefix đặc thù
            pass 
        elif text.startswith("Địa chỉ") and "BÊN B" not in text:
             # Đoạn text của template mẫu thường là "Địa chỉ\t\t: ..."
             # Ta sẽ dùng regex hoặc kiểm tra độ gần với block Bên B
             pass

        # Giải pháp an toàn hơn: Thay thế dựa trên tiền tố đặc thù trong template
        if text.startswith("Tài khoản số"):
            _rewrite_paragraph_text(paragraph, CONTRACT_SELLER["line2_account"])
        elif text.startswith("Địa chỉ") and "Huỳnh Đình Hai" in text: # Của Việt Nhật
            _rewrite_paragraph_text(paragraph, CONTRACT_SELLER["line3_address"])
        elif text.startswith("Điện thoại") and "086 224" in text:
            _rewrite_paragraph_text(paragraph, CONTRACT_SELLER["line4_phone"])
        elif text.startswith("Mã số thuế") and "0316419340" in text:
            _rewrite_paragraph_text(paragraph, CONTRACT_SELLER["line5_tax_code"])
        elif text.startswith("Đại diện là") and "NGÔ QUANG KHẢI" in text:
            _rewrite_paragraph_text(paragraph, CONTRACT_SELLER["line6_rep"])
        
        # Các mục điều khoản khác
        elif text.startswith("- Đợt 1:") and "{{tong_cong}}" not in text:
            _rewrite_paragraph_text(paragraph, CONTRACT_PAYMENT_TERMS)
        elif text.startswith("Địa điểm giao hàng:") and "{{dia_chi}}" not in text:
            _rewrite_paragraph_text(paragraph, CONTRACT_DELIVERY_LOCATION)


def _validate_template_schema(document, doc_type):
    content = "\n".join(
        paragraph.text for paragraph in document.paragraphs
    ) + "\n" + "\n".join(
        paragraph.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )

    if doc_type == "quotation":
        required_tokens = [
            "{{TEN_DOI_TAC}}",
            "{{dia_chi}}",
            "{{NGAY_HIEN_TAI}}",
            "{{tong_cong}}",
            "{%tr for item in items %}",
            "{%tr endfor %}",
        ]
    else:
        required_tokens = [
            "{{TEN_DOI_TAC}}",
            "{{dia_chi}}",
            "{{NGUOI_DAI_DIEN}}",
            "{{tong_cong}}",
            "{{tien_bang_chu}}",
            "{%tr for item in items %}",
            "{%tr endfor %}",
        ]

    missing_tokens = [token for token in required_tokens if token not in content]
    if missing_tokens:
        raise ValueError(
            "Template thiếu placeholder/tag bắt buộc sau khi chuẩn hóa: "
            + ", ".join(missing_tokens)
        )


def sanitize_template(input_path, output_path, doc_type):
    try:
        document = Document(input_path)

        if doc_type == "quotation":
            _normalize_quotation_template(document)
        else:
            _normalize_contract_template(document)

        for table in document.tables:
            for row in table.rows:
                _normalize_inline_loop_row(table, row)

        for table in document.tables:
            for row in table.rows:
                row_text = "".join(cell.text for cell in row.cells)
                if re.search(r"{%\s*tr\b", row_text) and "{{" not in row_text:
                    cells = row.cells
                    if len(cells) > 1:
                        main_cell = cells[0]
                        for other_cell in cells[1:]:
                            main_cell.merge(other_cell)

        for paragraph in document.paragraphs:
            _clean_p_final(paragraph)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _clean_p_final(paragraph)

        _validate_template_schema(document, doc_type)
        document.save(output_path)
        return True
    except Exception as err:
        print(f"Lỗi khi vệ sinh template: {err}")
        return False


def validate_payload(data):
    if not isinstance(data, dict):
        raise ValueError("JSON đầu vào phải là object.")

    doc_type = str(data.get("type", "contract")).strip().lower()
    if doc_type not in REQUIRED_FIELDS:
        raise ValueError("`type` chỉ được phép là `contract` hoặc `quotation`.")

    missing_fields = [field for field in REQUIRED_FIELDS[doc_type] if not data.get(field)]
    if missing_fields:
        raise ValueError(f"Thiếu trường bắt buộc: {', '.join(missing_fields)}.")

    if doc_type == "quotation":
        items = data.get("items")
        if not isinstance(items, list):
            raise ValueError("`items` phải là danh sách.")
        if not items:
            raise ValueError("`items` không được để trống.")
        for index, item in enumerate(items):
            if not isinstance(item, dict):
                raise ValueError(f"`items[{index}]` phải là object.")
            for key in ("model", "so_luong", "don_gia"):
                if item.get(key) in (None, ""):
                    raise ValueError(f"Thiếu `{key}` trong `items[{index}]`.")

    return doc_type


def enrich_quotation_data(data, csv_path):
    total = 0.0
    for index, item in enumerate(data.get("items", []), start=1):
        item["stt"] = index
        item["model"] = str(item.get("model", "")).strip()

        product_info = lookup_product(item["model"], csv_path)
        item["san_pham"] = product_info.get("san_pham", "") if product_info else ""
        item["dung_tich"] = product_info.get("dung_tich", "") if product_info else ""
        item["kich_thuoc"] = product_info.get("kich_thuoc", "") if product_info else ""
        item.setdefault("dvt", "Cái")

        quantity = force_float(item.get("so_luong", 0))
        price = force_float(item.get("don_gia", 0))
        line_total = quantity * price

        item["so_luong"] = int(quantity) if quantity.is_integer() else quantity
        item["don_gia_fmt"] = format_currency(price)
        item["thanh_tien"] = format_currency(line_total)
        total += line_total

    data["tong_cong"] = format_currency(total)
    data.setdefault("dia_chi", "")
    data.setdefault("NGAY_HIEN_TAI", datetime.now().strftime("%d/%m/%Y"))


def enrich_contract_data(data):
    for key in ("sdt", "stk", "mst", "chuc_vu", "tong_cong", "tien_bang_chu"):
        data.setdefault(key, "")
    data.setdefault("items", [])




def process_document_data(data, base_dir):
    output_dir = os.path.join(base_dir, "..", "output")
    csv_path = os.path.join(base_dir, "..", "assets", "product_db.csv")

    try:
        doc_type = validate_payload(data)
    except ValueError as err:
        print(f"Lỗi dữ liệu đầu vào: {err}")
        return {"success": False}

    safe_name = sanitize_filename(data.get("TEN_DOI_TAC", "Khach"))
    current_date = datetime.now().strftime("%Y%m%d")

    if doc_type == "quotation":
        template_path = os.path.join(base_dir, "..", "assets", "template_bao_gia.docx")
        file_name = f"Bao_gia_{safe_name}_{current_date}.docx"
        enrich_quotation_data(data, csv_path)
    else:
        template_path = os.path.join(base_dir, "..", "assets", "template_hop_dong.docx")
        file_name = f"Hop_dong_{safe_name}.docx"
        enrich_contract_data(data)

    os.makedirs(output_dir, exist_ok=True)
    fd, temp_template = tempfile.mkstemp(
        prefix="temp_tpl_",
        suffix=".docx",
        dir=tempfile.gettempdir(),
    )
    os.close(fd)

    try:
        if not sanitize_template(template_path, temp_template, doc_type):
            return {"success": False}

        doc = DocxTemplate(temp_template)
        doc.render(data)
        output_path = os.path.join(output_dir, file_name)
        doc.save(output_path)
        print(f"THANH CONG: {file_name}")

        return {
            "success": True,
            "doc_type": doc_type,
            "output_path": output_path,
        }
    except Exception as err:
        print(f"Lỗi render: {err}")
        return {"success": False}
    finally:
        if os.path.exists(temp_template):
            os.remove(temp_template)


def process_document(json_file_path):
    base_dir = os.path.dirname(os.path.abspath(__file__))
    try:
        with open(json_file_path, "r", encoding="utf-8") as json_file:
            data = json.load(json_file)
    except (OSError, json.JSONDecodeError) as err:
        print(f"Lỗi đọc JSON: {err}")
        return False

    result = process_document_data(data, base_dir)
    return bool(result.get("success"))


def parse_args(argv):
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--validate-models",
        action="store_true",
        help="Kiểm tra tính hợp lệ của các model trong temp_data.json dựa trên database.",
    )
    return parser.parse_args(argv)


def main(argv=None):
    args = parse_args(argv)
    base_dir = os.path.dirname(os.path.abspath(__file__))

    json_path = os.path.join(base_dir, "temp_data.json")
    if len(args) > 0 and not isinstance(args, argparse.Namespace):
        pass # Handle if we want to pass argv[1] directly
    if not os.path.exists(json_path) and len(sys.argv) > 1 and sys.argv[1].endswith(".json"):
        json_path = sys.argv[1]

    if not os.path.exists(json_path):
        print("Không tìm thấy temp_data.json.")
        sys.exit(1)

    # Logic kiểm tra model nếu được yêu cầu hoặc khi chạy bình thường
    csv_path = os.path.join(base_dir, "..", "assets", "product_db.csv")
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        items = data.get("items", [])
        invalid_found = False
        for item in items:
            model = item.get("model")
            if not model:
                continue
            if not lookup_product(model, csv_path):
                suggestions = find_closest_models(model, csv_path)
                print(f"CANH BAO: Model '{model}' khong tim thay.")
                if suggestions:
                    print(f"  -> Suggestions: {', '.join(suggestions)}")
                invalid_found = True
        
        if args.validate_models:
            sys.exit(1 if invalid_found else 0)
            
    except Exception as err:
        print(f"Lỗi khi kiểm tra model: {err}")

    success = process_document(json_path)
    if success:
        try:
            os.remove(json_path)
        except OSError:
            pass
        sys.exit(0)

    print("Giữ lại temp_data.json để kiểm tra lỗi.")
    sys.exit(1)


if __name__ == "__main__":
    main(sys.argv[1:])
